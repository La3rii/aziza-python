import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import random
import string

def generate_immatriculation():
    letters = ''.join(random.choices(string.ascii_uppercase, k=3))
    numbers = ''.join(random.choices(string.digits, k=4))
    return f"{letters}{numbers}"

def process_files(file_path, session_id, processed_folder):
    try:
        # Read and process CSV
        df = pd.read_csv(file_path)
        df["Salaire"] = df["Nombre de jours"] * 3
        df = df[(df["Salaire"] >= 0) & (df["Salaire"] <= 27)]
        df = df.sort_values(by="Salaire")
        
        # Generate immatriculation for each row
        df["Immatriculation"] = df.apply(lambda x: generate_immatriculation(), axis=1)
        
        # Save processed CSV
        csv1_path = os.path.join(processed_folder, f'{session_id}_personnes_100_salaire.csv')
        df.to_csv(csv1_path, index=False, encoding="utf-8")
        
        # Generate grouped CSV
        grouped_csv_path = os.path.join(processed_folder, f'{session_id}_regroupement_par_salaire.csv')
        generate_grouped_csv(df, grouped_csv_path)
        
        # Generate Word document
        doc_path = os.path.join(processed_folder, f'{session_id}_regroupement_par_salaire.docx')
        generate_word_document(df, doc_path)
        
        # Generate Excel file
        excel_path = os.path.join(processed_folder, f'{session_id}_regroupement_par_salaire.xlsx')
        generate_excel_file(df, excel_path)
        
        return True
        
    except Exception as e:
        raise e

def generate_grouped_csv(df, output_path):
    output_rows = []
    group_number = 1

    for salaire, group in df.groupby("Salaire"):
        nombre_personnes = len(group)
        nombre_jours = group["Nombre de jours"].iloc[0]
        
        output_rows.append([f"GROUPE {group_number} - Salaire: {salaire}DT"])
        output_rows.append([f"Nombre de personnes: {nombre_personnes}"])
        output_rows.append([f"Nombre de jours: {nombre_jours}"])
        output_rows.append([])
        
        output_rows.append([
            "Nom", "Prénom", "Immatriculation", "ID Groupe", "Salaire (DT)", "Jours travaillés"
        ])
        
        for _, row in group.iterrows():
            output_rows.append([
                row["Nom"], 
                row["Prénom"],
                row["Immatriculation"],
                f"G{group_number}", 
                salaire, 
                nombre_jours
            ])
        
        output_rows.append([])
        output_rows.append(["="*50])
        output_rows.append([])
        
        group_number += 1

    pd.DataFrame(output_rows).to_csv(
        output_path, 
        index=False, 
        header=False, 
        encoding="utf-8"
    )

def generate_word_document(df, output_path):
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    title = doc.add_paragraph("Rapport de Groupement par Salaire")
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)
    doc.add_paragraph("\n")

    group_number = 1

    for salaire, group in df.groupby("Salaire"):
        nombre_personnes = len(group)
        nombre_jours = group["Nombre de jours"].iloc[0]
        
        header = doc.add_paragraph()
        header.add_run(f"GROUPE {group_number}").bold = True
        header.add_run(" - ")
        header.add_run(f"Salaire: {salaire}DT").bold = True
        
        info = doc.add_paragraph()
        info.add_run(f"Nombre de personnes: {nombre_personnes}\t")
        info.add_run(f"Jours travaillés: {nombre_jours}")
        
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        headers = ["Nom", "Prénom", "Immatriculation", "ID Groupe", "Salaire (DT)", "Jours"]
        for i, header_text in enumerate(headers):
            hdr_cells[i].text = header_text
            hdr_cells[i].paragraphs[0].runs[0].bold = True
        
        for _, row in group.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = row["Nom"]
            row_cells[1].text = row["Prénom"]
            row_cells[2].text = row["Immatriculation"]
            row_cells[3].text = f"G{group_number}"
            row_cells[4].text = str(salaire)
            row_cells[5].text = str(nombre_jours)
        
        doc.add_paragraph("\n")
        doc.add_paragraph("-" * 50)
        doc.add_paragraph("\n")
        
        group_number += 1

    doc.add_paragraph("\n")
    footer = doc.add_paragraph(f"Total des groupes: {group_number-1}")
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    footer.runs[0].italic = True

    doc.save(output_path)

def generate_excel_file(df, output_path):
    excel_data = []
    group_number = 1

    for salaire, group in df.groupby("Salaire"):
        nombre_personnes = len(group)
        nombre_jours = group["Nombre de jours"].iloc[0]

        excel_data.append({"A": f"GROUPE {group_number} - Salaire: {salaire}DT"})
        excel_data.append({"A": f"Nombre de personnes: {nombre_personnes}"})
        excel_data.append({"A": f"Nombre de jours: {nombre_jours}"})
        excel_data.append({})  # Empty row
        
        # Header row
        excel_data.append({
            "A": "Nom", 
            "B": "Prénom", 
            "C": "Immatriculation", 
            "D": "ID Groupe", 
            "E": "Salaire (DT)", 
            "F": "Jours"
        })
        
        for _, row in group.iterrows():
            excel_data.append({
                "A": row["Nom"],
                "B": row["Prénom"],
                "C": row["Immatriculation"],
                "D": f"G{group_number}",
                "E": salaire,
                "F": nombre_jours
            })
        
        excel_data.append({})  # Empty row
        group_number += 1
    
    # Convert to DataFrame and save
    excel_df = pd.DataFrame(excel_data)
    excel_df.to_excel(output_path, index=False, engine='openpyxl')